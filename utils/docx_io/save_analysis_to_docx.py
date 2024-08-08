from docx import Document
import logging

logging.basicConfig(level=logging.INFO)

def add_heading(doc: Document, text: str, level: int) -> None:
    """Adds a heading to the document.
    
    Args:
        doc (Document): The Word document.
        text (str): The heading text.
        level (int): The heading level.
    """
    doc.add_heading(text.strip(), level=level)

def add_bold_paragraph(doc: Document, line: str) -> None:
    """Adds a paragraph with bold text to the document.
    
    Args:
        doc (Document): The Word document.
        line (str): The line containing bold and normal text.
    """
    try:
        bold_text, normal_text = line.split('**')[1], line.split('**')[2].strip()
        p = doc.add_paragraph()
        p.add_run(bold_text + ':').bold = True
        p.add_run(' ' + normal_text)
    except IndexError:
        logging.error(f"Error processing bold text in line: {line}")

def save_text_to_docx(text_data: str, output_path: str) -> None:
    """Saves text data to a Word document.
    
    Args:
        text_data (str): A string containing the entire text data.
        output_path (str): Path to the output Word document.
    """
    doc = Document()
    lines = text_data.split('\n')

    for line in lines:
        try:
            if line.startswith('####'):
                add_heading(doc, line.replace('####', ''), level=1)
            elif line.startswith('###'):
                add_heading(doc, line.replace('###', ''), level=2)
            elif line.startswith('- **'):
                add_bold_paragraph(doc, line)
            elif line.strip() == '':
                continue
            else:
                doc.add_paragraph(line.strip())
        except Exception as e:
            logging.error(f"Error processing line: {line}. Error: {e}")

    try:
        doc.save(output_path)
        logging.info(f"Document saved successfully at {output_path}")
    except Exception as e:
        logging.error(f"Error saving document at {output_path}: {e}")
